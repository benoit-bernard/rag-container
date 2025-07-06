"""
Metadata Manager Service
"""
import logging
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path
from app.core.config import config

logger = logging.getLogger(__name__)

class MetadataManager:
    def __init__(self):
        self.validation_strict = config.METADATA_VALIDATION_STRICT
        self.supported_extensions = config.SUPPORTED_EXTENSIONS
        logger.debug(f"MetadataManager initialized (strict: {self.validation_strict})")
    
    def validate_metadata(self, metadata: Dict[str, Any]) -> bool:
        """
        Validate metadata according to configuration
        Returns True if valid, False otherwise
        """
        try:
            if not self.validation_strict:
                return True  # Permissive mode
            
            # Required fields for strict validation
            required_fields = ["size", "modified", "extension", "name"]
            
            for field in required_fields:
                if field not in metadata:
                    logger.warning(f"Missing required field in metadata: {field}")
                    return False
            
            # Validate field types
            if not isinstance(metadata.get("size"), int):
                logger.warning("Size field must be integer")
                return False
            
            if not isinstance(metadata.get("name"), str):
                logger.warning("Name field must be string")
                return False
            
            # Validate extension
            extension = metadata.get("extension", "").lower()
            if extension not in self.supported_extensions:
                logger.warning(f"Unsupported extension: {extension}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating metadata: {e}")
            return False
    
    def enrich_metadata(self, metadata: Dict[str, Any], file_path: Optional[Path] = None) -> Dict[str, Any]:
        """
        Enrich metadata with additional information
        """
        try:
            enriched = metadata.copy()
            
            # Add validation timestamp
            enriched["validation_timestamp"] = datetime.now().isoformat()
            enriched["validation_strict"] = self.validation_strict
            
            # Add configuration context
            enriched["metadata_config"] = {
                "validation_mode": "strict" if self.validation_strict else "permissive",
                "supported_extensions": self.supported_extensions
            }
            
            # Add file categorization
            extension = metadata.get("extension", "").lower()
            enriched["file_category"] = self._categorize_file(extension)
            
            # Add size category
            size = metadata.get("size", 0)
            enriched["size_category"] = self._categorize_size(size)
            
            # Add additional file path info if provided
            if file_path:
                enriched["path_info"] = {
                    "depth": len(file_path.parts) - len(config.DOCUMENTS_DIR.parts) - 1,
                    "parent_directory": file_path.parent.name,
                    "is_in_subdirectory": len(file_path.parts) > len(config.DOCUMENTS_DIR.parts) + 1
                }
            
            return enriched
            
        except Exception as e:
            logger.error(f"Error enriching metadata: {e}")
            return metadata
    
    def _categorize_file(self, extension: str) -> str:
        """Categorize file by extension"""
        categories = {
            "text": [".txt", ".md"],
            "document": [".pdf", ".docx", ".doc", "ppt", ".pptx"],
            "code": [".py", ".c", ".cpp", ".cs", ".js", ".ts"],
            "data": [".json", ".csv", ".xml"],
            "other": []
        }
        
        for category, extensions in categories.items():
            if extension in extensions:
                return category
        
        return "other"
    
    def _categorize_size(self, size_bytes: int) -> str:
        """Categorize file by size"""
        if size_bytes < 1024:  # < 1KB
            return "tiny"
        elif size_bytes < 1024 * 1024:  # < 1MB
            return "small"
        elif size_bytes < 10 * 1024 * 1024:  # < 10MB
            return "medium"
        elif size_bytes < 100 * 1024 * 1024:  # < 100MB
            return "large"
        else:
            return "very_large"
    
    def extract_content_metadata(self, file_path: Path) -> Dict[str, Any]:
        try:
            extension = file_path.suffix.lower()
            
            if extension in [".txt", ".md"]:
                return self._extract_text_content_metadata(file_path)
            elif extension == ".json":
                return self._extract_json_content_metadata(file_path)
            elif extension == ".py":
                return self._extract_python_content_metadata(file_path)
            elif extension in [".c", ".cpp", ".cs", ".js", ".ts"] :
                return self._extract_code_content_metadata(file_path)
            elif extension in [".pdf", ".docx", ".doc", ".ppt", ".pptx"]:
                return self._extract_document_content_metadata(file_path)
            else:
                return {"content_type": "unsupported"}
                
        except Exception as e:
            logger.warning(f"Error extracting content metadata from {file_path}: {e}")
            return {"content_error": str(e)}
        
    def _extract_document_content_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from document files (PDF, DOCX, etc.)"""
        try:
            from PyPDF2 import PdfReader
            from docx import Document
            
            if file_path.suffix.lower() == ".pdf":
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    metadata = {
                        "content_type": "pdf",
                        "page_count": len(reader.pages),
                        "author": reader.metadata.get('/Author', 'Unknown'),
                        "title": reader.metadata.get('/Title', 'Unknown'),
                        "subject": reader.metadata.get('/Subject', 'Unknown')
                    }
                return metadata
            
            elif file_path.suffix.lower() in [".docx", ".doc"]:
                doc = Document(file_path)
                metadata = {
                    "content_type": "docx",
                    "paragraph_count": len(doc.paragraphs),
                    "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
                    "author": doc.core_properties.author or 'Unknown',
                    "title": doc.core_properties.title or 'Unknown'
                }
                return metadata
            
            else:
                return {"content_type": "unsupported"}
                
        except Exception as e:
            logger.warning(f"Error extracting document metadata: {e}")
            return {"content_error": str(e)}
    
    def _extract_text_content_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.splitlines()
            words = content.split()
            
            metadata = {
                "content_type": "text",
                "line_count": len(lines),
                "word_count": len(words),
                "character_count": len(content),
                "non_empty_lines": len([line for line in lines if line.strip()]),
                "encoding": "utf-8"
            }
            
            # Add specific markdown metadata
            if file_path.suffix.lower() == ".md":
                metadata.update({
                    "markdown_headers": content.count('#'),
                    "markdown_links": content.count('['),
                    "markdown_code_blocks": content.count('```')
                })
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting text metadata: {e}")
            return {"content_error": str(e)}
    
    def _extract_json_content_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from JSON files"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            metadata = {
                "content_type": "json",
                "json_type": type(data).__name__,
                "valid_json": True
            }
            
            if isinstance(data, dict):
                metadata.update({
                    "key_count": len(data.keys()),
                    "keys": list(data.keys())[:10],  # First 10 keys
                    "nested_levels": self._count_json_depth(data)
                })
            elif isinstance(data, list):
                metadata.update({
                    "array_length": len(data),
                    "item_types": list(set(type(item).__name__ for item in data[:10]))
                })
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting JSON metadata: {e}")
            return {"content_error": str(e), "valid_json": False}
    
    def _extract_python_content_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from Python files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.splitlines()
            
            metadata = {
                "content_type": "python",
                "line_count": len(lines),
                "import_count": len([line for line in lines if line.strip().startswith(('import ', 'from '))]),
                "function_count": content.count('def '),
                "class_count": content.count('class '),
                "comment_lines": len([line for line in lines if line.strip().startswith('#')]),
                "docstring_count": content.count('"""') // 2,
                "encoding": "utf-8"
            }
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting Python metadata: {e}")
            return {"content_error": str(e)}

    def _extract_code_content_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from code files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.splitlines()
            
            metadata = {
                "content_type": file_path.suffix.lower(),
                "line_count": len(lines),
                "import_count": len([line for line in lines if line.strip().startswith(('import ', 'from '))]),
                "function_count": content.count('def '),
                "class_count": content.count('class '),
                "comment_lines": len([line for line in lines if line.strip().startswith('#')]),
                "docstring_count": content.count('"""') // 2,
                "encoding": "utf-8"
            }
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting Python metadata: {e}")
            return {"content_error": str(e)}
    
    def _count_json_depth(self, obj, depth=0) -> int:
        """Count maximum nesting depth in JSON object"""
        if isinstance(obj, dict):
            if not obj:
                return depth
            return max(self._count_json_depth(value, depth + 1) for value in obj.values())
        elif isinstance(obj, list):
            if not obj:
                return depth
            return max(self._count_json_depth(item, depth + 1) for item in obj)
        else:
            return depth
    
    def get_validation_info(self) -> Dict[str, Any]:
        """Get information about current validation configuration"""
        return {
            "validation_strict": self.validation_strict,
            "supported_extensions": self.supported_extensions,
            "required_fields": ["size", "modified", "extension", "name"] if self.validation_strict else [],
            "configuration_source": "centralized config"
        }